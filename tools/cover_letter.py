"""
Tool: cover_letter.py

Purpose:
    Creates cover letter .docx files and uploads them to Google Drive.
    Cover letters are stored in a "Cover Letters" subfolder inside the
    project Drive folder.

Actions:
    create_cover_letter
        Takes a plain-text cover letter from a file, creates a formatted
        .docx, and uploads it to Google Drive. Returns the shareable URL.

Usage:
    python tools/cover_letter.py \
        --company "Acme Corp" \
        --job_title "Engineering Manager" \
        --content_file ".tmp/cover_letter_draft.txt"

Returns:
    Google Drive shareable URL printed to stdout.

Exit codes:
    0 = success
    1 = error (details on stderr)
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.google_auth import get_credentials
from tools.drive_helpers import get_or_create_folder

TMP_FOLDER = Path(__file__).parent.parent / ".tmp"
ENV_PATH = Path(__file__).parent.parent / ".env"
COVER_LETTER_SUBFOLDER = "Cover Letters"


def _read_env_value(key: str) -> str:
    if not ENV_PATH.exists():
        return ""
    for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line.startswith(f"{key}="):
            return line.split("=", 1)[1].strip()
    return ""


def sanitize_filename(text: str) -> str:
    sanitized = re.sub(r"[^\w\s-]", "", text)
    sanitized = re.sub(r"[\s_]+", "-", sanitized)
    return sanitized.strip("-")


def create_cover_letter(company: str, job_title: str, content_file: Path) -> str:
    """
    Create a formatted cover letter .docx and upload to Google Drive.

    Args:
        company:      Company name (used in filename)
        job_title:    Job title (used in filename)
        content_file: Path to plain text cover letter content

    Returns:
        Google Drive shareable URL
    """
    text = content_file.read_text(encoding="utf-8").strip()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Write paragraphs
    for line in text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line.strip())

    # Save locally
    TMP_FOLDER.mkdir(exist_ok=True)
    safe_company = sanitize_filename(company)
    safe_title = sanitize_filename(job_title)
    filename = f"{safe_company}-{safe_title}-Cover-Letter.docx"
    tmp_path = TMP_FOLDER / filename

    doc.save(str(tmp_path))

    # Upload to Drive
    creds = get_credentials()
    service = build("drive", "v3", credentials=creds)

    project_folder_id = _read_env_value("DRIVE_FOLDER_ID") or None
    cover_folder_id = get_or_create_folder(
        service, COVER_LETTER_SUBFOLDER, parent_id=project_folder_id
    )

    media = MediaFileUpload(
        str(tmp_path),
        mimetype="application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document",
        resumable=True,
    )
    result = service.files().create(
        body={"name": filename, "parents": [cover_folder_id]},
        media_body=media,
        fields="id,webViewLink"
    ).execute()
    service.permissions().create(
        fileId=result["id"],
        body={"type": "anyone", "role": "reader"},
    ).execute()
    return result["webViewLink"]


def main():
    parser = argparse.ArgumentParser(description="Cover letter creation tool.")
    parser.add_argument("--company", required=True)
    parser.add_argument("--job_title", required=True)
    parser.add_argument("--content_file", required=True,
                        help="Path to plain text cover letter content.")
    args = parser.parse_args()

    content_path = Path(args.content_file)
    if not content_path.exists():
        print(f"ERROR: content_file not found: {content_path}", file=sys.stderr)
        sys.exit(1)

    try:
        url = create_cover_letter(args.company, args.job_title, content_path)
        print(url)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
