"""
Tool: tailor_resume.py

Purpose:
    Template-based resume tailoring. Reads the original .docx structure,
    lets Claude rewrite the text content, then rebuilds the document using
    the original as a template — preserving all fonts, sizes, spacing,
    bold/italic, colors, and layout exactly.

    The actual rewriting logic lives in the agent (Claude), not here.
    This tool handles all file mechanics and formatting preservation.

Actions:
    read_resume_structured
        Reads the resume from resume/ and returns a numbered paragraph list
        with formatting metadata. Claude uses this to rewrite text while
        knowing exactly which paragraphs are headers, bullets, etc.

    create_doc_from_template
        Takes Claude's rewritten paragraph list, applies it back onto a copy
        of the original .docx, and uploads the result to Google Drive.
        Preserves ALL original formatting — fonts, sizes, colors, spacing.
        ALWAYS use --content_file. Never pass content inline (shell will
        corrupt $100K → 00K and similar).

    convert_to_docx
        Converts a .doc file to .docx using LibreOffice (if installed).
        Run this once if the resume is in the older .doc format.

Usage:
    python tools/tailor_resume.py --action read_resume_structured

    python tools/tailor_resume.py --action create_doc_from_template \\
        --company "Acme Corp" \\
        --job_title "Software Engineer" \\
        --content_file ".tmp/resume_draft.txt"

    python tools/tailor_resume.py --action convert_to_docx

Content file format (for create_doc_from_template):
    One paragraph per line, prefixed with its number and a pipe:
        1|Jordan Lee
        2|City, State | 000-000-0000 | candidate@example.com
        3|SUMMARY
        4|Rewritten summary text here...
    Blank lines are preserved as spacers. Empty paragraphs (headers/spacers
    with no text) are carried over from the original unchanged.

Returns:
    read_resume_structured  → numbered paragraph list printed to stdout
    create_doc_from_template → Google Drive URL printed to stdout
    convert_to_docx          → path to converted .docx file

Exit codes:
    0 = success
    1 = error (details on stderr)
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.google_auth import get_credentials
from tools.drive_helpers import get_or_create_folder

RESUME_FOLDER = Path(__file__).parent.parent / "resume"
TMP_FOLDER = Path(__file__).parent.parent / ".tmp"
TAILORED_SUBFOLDER_NAME = "Tailored Resumes"
ENV_PATH = Path(__file__).parent.parent / ".env"


def _read_env_value(key: str) -> str:
    """Read a single value from .env. Returns empty string if not found."""
    if not ENV_PATH.exists():
        return ""
    for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line.startswith(f"{key}="):
            return line.split("=", 1)[1].strip()
    return ""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def sanitize_filename(text: str) -> str:
    """Convert a string to a safe filename component."""
    sanitized = re.sub(r"[^\w\s-]", "", text)
    sanitized = re.sub(r"[\s_]+", "-", sanitized)
    return sanitized.strip("-")


def find_resume_file() -> Path:
    """
    Find the resume file in the resume/ folder.
    Prefers .docx over .doc over .pdf.
    Raises FileNotFoundError if nothing is found.
    """
    for ext in ("*.docx", "*.doc", "*.pdf"):
        matches = [f for f in RESUME_FOLDER.glob(ext)
                   if f.name.lower() != "readme.txt"]
        if matches:
            return matches[0]
    raise FileNotFoundError(
        f"No resume file found in {RESUME_FOLDER}. "
        "Add a .docx, .doc, or .pdf file to the resume/ folder."
    )


def upload_to_drive(local_path: Path, filename: str) -> str:
    """
    Upload a file to the 'Tailored Resumes' subfolder in the project Drive folder.
    If DRIVE_FOLDER_ID is set in .env, creates 'Tailored Resumes' inside it.
    Otherwise falls back to Drive root.
    Returns the shareable URL.
    """
    creds = get_credentials()
    service = build("drive", "v3", credentials=creds)

    project_folder_id = _read_env_value("DRIVE_FOLDER_ID") or None
    tailored_folder_id = get_or_create_folder(
        service, TAILORED_SUBFOLDER_NAME, parent_id=project_folder_id
    )

    media = MediaFileUpload(
        str(local_path),
        mimetype="application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document",
        resumable=True,
    )
    result = service.files().create(
        body={"name": filename, "parents": [tailored_folder_id]},
        media_body=media,
        fields="id,webViewLink"
    ).execute()
    service.permissions().create(
        fileId=result["id"],
        body={"type": "anyone", "role": "reader"},
    ).execute()
    return result["webViewLink"]


# ---------------------------------------------------------------------------
# Actions
# ---------------------------------------------------------------------------

def read_resume_structured() -> str:
    """
    Read the resume .docx and return a numbered paragraph list with
    formatting notes so Claude knows the document structure.

    Output format (one paragraph per line):
        1|[HEADER] Jordan Lee
        2|[NORMAL] City, State | candidate@example.com
        3|[EMPTY]
        4|[HEADER] PROFESSIONAL SUMMARY
        5|[NORMAL] Accomplished lease accounting manager...

    Tags:
        [HEADER]  — bold, heading style, or short ALL-CAPS line
        [BULLET]  — list item
        [NORMAL]  — regular body text
        [EMPTY]   — blank paragraph (spacer — do not rewrite)

    Claude should rewrite [NORMAL] and [BULLET] paragraphs only.
    [HEADER] paragraphs may be left as-is or lightly adjusted.
    [EMPTY] paragraphs must be returned unchanged.

    Returns:
        Multi-line string — one numbered paragraph per line.

    Raises:
        FileNotFoundError if no resume file found
        ValueError if file is .doc — must be converted first
    """
    resume_path = find_resume_file()

    if resume_path.suffix.lower() == ".doc":
        raise ValueError(
            f"Resume is a .doc file ({resume_path.name}). "
            "Run: python tools/tailor_resume.py --action convert_to_docx\n"
            "Then retry. (Or open the file in Word and Save As .docx)"
        )

    doc = Document(str(resume_path))
    lines = []

    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()

        if not text:
            lines.append(f"{i}|[EMPTY]")
            continue

        # Detect paragraph type
        style_name = para.style.name.lower() if para.style else ""
        is_heading = "heading" in style_name
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        is_list = "list" in style_name or text.startswith(("•", "-", "–", "*"))
        is_allcaps = text.isupper() and len(text) > 2

        if is_heading or is_bold or is_allcaps:
            tag = "[HEADER]"
        elif is_list:
            tag = "[BULLET]"
        else:
            tag = "[NORMAL]"

        lines.append(f"{i}|{tag} {text}")

    return "\n".join(lines)


def convert_to_docx() -> str:
    """
    Convert a .doc file in the resume/ folder to .docx using LibreOffice.

    Requires LibreOffice to be installed. On Windows, installs via winget
    if not found.

    Returns:
        Path string of the converted .docx file.

    Raises:
        FileNotFoundError if no .doc file is found
        RuntimeError if conversion fails
    """
    doc_files = list(RESUME_FOLDER.glob("*.doc"))
    if not doc_files:
        raise FileNotFoundError("No .doc file found in resume/ folder.")

    doc_path = doc_files[0]

    # Try LibreOffice
    soffice_candidates = [
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]

    soffice = None
    for candidate in soffice_candidates:
        try:
            result = subprocess.run(
                [candidate, "--version"],
                capture_output=True, timeout=5
            )
            if result.returncode == 0:
                soffice = candidate
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    if not soffice:
        raise RuntimeError(
            "LibreOffice is not installed. To convert your .doc file:\n"
            "  Option 1: Install LibreOffice (free): https://www.libreoffice.org\n"
            "  Option 2: Open the file in Microsoft Word → File → Save As → .docx\n"
            "Then place the .docx file in the resume/ folder and remove the .doc file."
        )

    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", str(doc_path),
         "--outdir", str(RESUME_FOLDER)],
        capture_output=True, text=True, timeout=30
    )

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed:\n{result.stderr}")

    converted = doc_path.with_suffix(".docx")
    if not converted.exists():
        raise RuntimeError("Conversion appeared to succeed but .docx was not created.")

    return str(converted)


def create_doc_from_template(company: str, job_title: str, content_file: Path) -> str:
    """
    Build a tailored resume by applying rewritten text onto the original
    .docx template, preserving all formatting exactly.

    The content_file must contain the numbered paragraph format output
    by read_resume_structured(), with Claude's rewrites applied:
        1|[HEADER] Jordan Lee
        2|[NORMAL] Rewritten summary here...
        3|[EMPTY]
        ...

    Steps:
        1. Parse content_file into a dict: paragraph_number → rewritten text
        2. Copy original .docx to .tmp/
        3. Walk every paragraph in the copy
        4. For paragraphs with rewritten text: replace text, preserve formatting
        5. Upload to Drive, return URL

    Args:
        company:      Company name (used in filename)
        job_title:    Job title (used in filename)
        content_file: Path to the rewritten paragraph file

    Returns:
        Google Drive shareable URL
    """
    resume_path = find_resume_file()
    if resume_path.suffix.lower() != ".docx":
        raise ValueError(
            "Template approach requires a .docx file. "
            "Run --action convert_to_docx first."
        )

    # Parse content file → {para_num: rewritten_text}
    rewrites = {}
    for line in content_file.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or "|" not in line:
            continue
        num_str, rest = line.split("|", 1)
        try:
            num = int(num_str.strip())
        except ValueError:
            continue
        # Strip tag prefix ([HEADER], [NORMAL], [BULLET], [EMPTY])
        text = re.sub(r"^\[(HEADER|NORMAL|BULLET|EMPTY)\]\s*", "", rest).strip()
        if rest.strip() != "[EMPTY]":
            rewrites[num] = text

    # Copy original to .tmp/
    TMP_FOLDER.mkdir(exist_ok=True)
    safe_company = sanitize_filename(company)
    safe_title = sanitize_filename(job_title)
    filename = f"{safe_company}-{safe_title}-Resume.docx"
    tmp_path = TMP_FOLDER / filename
    shutil.copy2(str(resume_path), str(tmp_path))

    # Open copy and apply rewrites
    doc = Document(str(tmp_path))

    for para_num, para in enumerate(doc.paragraphs, start=1):
        if para_num not in rewrites:
            continue
        new_text = rewrites[para_num]
        if not new_text:
            continue

        # Capture formatting from existing runs
        run_formats = []
        for run in para.runs:
            run_formats.append({
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size,
                "color": (run.font.color.rgb
                          if run.font.color and run.font.color.type
                          else None),
            })

        # Clear all run text
        for run in para.runs:
            run.text = ""

        # Write new text into first run, applying original run's formatting
        if para.runs:
            para.runs[0].text = new_text
        else:
            # No runs existed — add one with default formatting
            run = para.add_run(new_text)
            if run_formats:
                fmt = run_formats[0]
                run.bold = fmt["bold"]
                run.italic = fmt["italic"]
                run.underline = fmt["underline"]
                if fmt["font_name"]:
                    run.font.name = fmt["font_name"]
                if fmt["font_size"]:
                    run.font.size = fmt["font_size"]

    doc.save(str(tmp_path))

    return upload_to_drive(tmp_path, filename)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Template-based resume tailoring.")
    parser.add_argument("--action", required=True,
                        choices=["read_resume_structured",
                                 "create_doc_from_template",
                                 "convert_to_docx"])
    parser.add_argument("--company", default=None)
    parser.add_argument("--job_title", default=None)
    parser.add_argument("--content_file", default=None,
                        help="Path to rewritten paragraph file. "
                             "Never pass content inline — shell corrupts special chars.")
    args = parser.parse_args()

    try:
        if args.action == "read_resume_structured":
            print(read_resume_structured())

        elif args.action == "convert_to_docx":
            path = convert_to_docx()
            print(f"Converted: {path}")

        elif args.action == "create_doc_from_template":
            for field in ("company", "job_title", "content_file"):
                if not getattr(args, field):
                    print(f"ERROR: --{field} required", file=sys.stderr)
                    sys.exit(1)
            content_path = Path(args.content_file)
            if not content_path.exists():
                print(f"ERROR: content_file not found: {content_path}",
                      file=sys.stderr)
                sys.exit(1)
            url = create_doc_from_template(args.company, args.job_title, content_path)
            print(url)

        sys.exit(0)

    except (FileNotFoundError, ValueError, RuntimeError) as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
